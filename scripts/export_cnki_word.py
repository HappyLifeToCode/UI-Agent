#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""CNKI 文献汇总 Word 导出：data/cnki_<关键词>/papers.json → <关键词>_汇总表.docx。

报告结构：
    标题（关键词 + 来源 + 生成时间）
    一、文献总表（序号/标题/作者/期刊/时间/关键词；序号与标题为
        指向摘要附录的内部超链接，Ctrl+点击跳转）
    二、摘要全文（每篇一个小节，标题处为对应书签锚点）

字体统一：中文宋体正文 / 黑体标题，西文 Times New Roman。

用法：
    python scripts/export_cnki_word.py --task 芍药甘草汤
    python scripts/export_cnki_word.py --all        # data/ 下所有 cnki_* 任务
"""
from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = PROJECT_ROOT / "data"

# 全文统一字体（与 export_word.py 同一约定）
BODY_ASCII_FONT = "Times New Roman"
BODY_EASTASIA_FONT = "宋体"
HEADING_EASTASIA_FONT = "黑体"


def _set_style_font(style, ascii_font: str, eastasia_font: str, size=None):
    style.font.name = ascii_font
    rfonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rfonts.attrib.pop(qn(f"w:{attr}"), None)
    rfonts.set(qn("w:eastAsia"), eastasia_font)
    if size is not None:
        style.font.size = size


def _add_bookmark(paragraph, name: str, bid: int):
    """给段落加书签（内部跳转锚点），供 w:hyperlink w:anchor 引用。"""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_link(paragraph, text: str, anchor: str):
    """在段落里追加一个指向文档内书签的超链接（蓝色下划线样式）。"""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BODY_ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), BODY_EASTASIA_FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(rfonts)
    rpr.append(color)
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def generate_report(task_dir: Path) -> Path | None:
    """为单个 cnki 任务目录生成 <关键词>_汇总表.docx；papers.json 缺失返回 None。"""
    pj = task_dir / "papers.json"
    if not pj.exists():
        return None
    data = json.loads(pj.read_text(encoding="utf-8"))
    papers = data.get("papers") or []

    doc = Document()
    _set_style_font(doc.styles["Normal"], BODY_ASCII_FONT, BODY_EASTASIA_FONT, Pt(10.5))
    for name in ("Title", "Heading 1", "Heading 2"):
        _set_style_font(doc.styles[name], BODY_ASCII_FONT, HEADING_EASTASIA_FONT)

    doc.add_heading(f"“{data.get('keyword') or task_dir.name}”知网文献汇总表", level=0)
    doc.add_paragraph(
        f"数据来源：{data.get('source', '中国知网（CNKI）')}（{data.get('sort', '默认相关度排序')}）    "
        f"命中 {data.get('total_results', '-')} 条，采集 {data.get('collected', f'{len(papers)} 篇')}    "
        f"生成时间：{datetime.now():%Y-%m-%d %H:%M}")

    # ---- 一、文献总表（序号/标题为指向摘要附录的内部超链接） ----
    doc.add_heading("一、文献总表", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for i, t in enumerate(("序号", "标题", "作者", "期刊/来源", "发表时间", "关键词")):
        c = table.rows[0].cells[i]
        c.text = t
        c.paragraphs[0].runs[0].font.bold = True
    for p in papers:
        anchor = f"abs_{p['rank']:02d}"
        cells = table.add_row().cells
        cells[0].text = ""
        _add_internal_link(cells[0].paragraphs[0], str(p.get("rank")), anchor)
        cells[1].text = ""
        _add_internal_link(cells[1].paragraphs[0], p.get("title") or "(无标题)", anchor)
        cells[2].text = p.get("authors") or "-"
        cells[3].text = p.get("source") or "-"
        cells[4].text = p.get("date") or "-"
        cells[5].text = "；".join(p.get("keywords") or [])
        if p.get("note"):
            cells[1].paragraphs[0].add_run(f"（{p['note']}）")

    # ---- 二、摘要全文（每篇标题为书签锚点） ----
    doc.add_heading("二、摘要全文", level=1)
    for p in papers:
        h = doc.add_heading(f"#{p.get('rank')}  {p.get('title') or '(无标题)'}", level=2)
        _add_bookmark(h, f"abs_{p['rank']:02d}", int(p.get("rank") or 0))
        doc.add_paragraph(f"{p.get('authors') or '-'}｜{p.get('source') or '-'}｜{p.get('date') or '-'}")
        doc.add_paragraph(p.get("abstract") or "（知网未提供摘要）")

    out = task_dir / f"{data.get('keyword') or task_dir.name}_汇总表.docx"
    doc.save(str(out))
    return out


def main():
    parser = argparse.ArgumentParser(description="把 CNKI 采集结果导出为汇总 Word")
    parser.add_argument("--task", action="append", help="关键词（可多次指定）")
    parser.add_argument("--all", action="store_true", help="data/ 下所有有 papers.json 的 cnki_ 任务")
    args = parser.parse_args()

    if not args.task and not args.all:
        parser.error("请指定 --task <关键词> 或 --all")

    task_dirs = [DATA_DIR / f"cnki_{k}" for k in args.task or []]
    if args.all:
        task_dirs += sorted(p.parent for p in DATA_DIR.glob("cnki_*/papers.json"))

    ok, skip = 0, 0
    for td in dict.fromkeys(task_dirs):  # 去重保序
        out = generate_report(td)
        if out:
            print(f"[OK] {out}")
            ok += 1
        else:
            print(f"[跳过] {td.name}：没有 papers.json")
            skip += 1
    print(f"完成：生成 {ok} 份，跳过 {skip} 个")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()

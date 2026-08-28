#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Word 报告导出：把 data/<task_id>/ 的采集结果生成可交付的 docx。

用法：
    python scripts/export_word.py --task task_0001          # 单个任务
    python scripts/export_word.py --task task_0001 --task task_0002
    python scripts/export_word.py --all                     # data/ 下所有有 result.json 的任务

也被 scripts/run_tasks.py 在三段式管线合并后 import 调用（generate_report）。

报告结构：
    标题（学者姓名 + task_id + 生成时间）
    一、作者信息（单位/兴趣/总被引/h-index/i10-index/主页链接 + 主页整页截图）
    二、近五年论文总表（年内序号/标题/GS被引/核查状态，年份分隔行带篇数）
    三、逐篇核查详情（按年份分小节，年内按 GS 被引降序）
        每篇：全局被引排名、标题、年份、GS 被引、核查状态、S2/OpenAlex 数据、
        DOI、期刊、备注，下方嵌入对应截图证据。

字段兼容：recent_papers 同时认 s2_*（Semantic Scholar 链路）和
openalex_*（OpenAlex 链路）两套命名（如旧数据 task_0001 是 s2 字段）。
"""
from __future__ import annotations  # Path | None 注解在 Python 3.9 下需延迟求值

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = PROJECT_ROOT / "data"

SCREENSHOT_WIDTH = Cm(15)  # 正文宽度约 16cm，留一点边距

# 全文统一字体：西文/数字 Times New Roman，中文正文宋体、标题黑体。
# python-docx 的 font.name 只设 w:rFonts 的 ascii/hAnsi，中文字体必须
# 显式写 w:eastAsia，否则中文回落到模板默认字体，同一段中西文粗细不一。
BODY_ASCII_FONT = "Times New Roman"
BODY_EASTASIA_FONT = "宋体"
HEADING_EASTASIA_FONT = "黑体"


def _set_style_font(style, ascii_font: str, eastasia_font: str, size=None):
    """统一设置样式的西文与中文字体（eastAsia 必须走 XML 显式设置）。"""
    style.font.name = ascii_font
    rfonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    # 标题样式自带 asciiTheme/eastAsiaTheme 主题引用，删除以免与显式字体歧义
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rfonts.attrib.pop(qn(f"w:{attr}"), None)
    rfonts.set(qn("w:eastAsia"), eastasia_font)
    if size is not None:
        style.font.size = size


def _ext_fields(paper: dict):
    """返回 (来源名, url, citations)，兼容 s2_* / openalex_* 两套字段。"""
    if "openalex_id" in paper or "openalex_url" in paper:
        return "OpenAlex", paper.get("openalex_url"), paper.get("openalex_citations")
    return "Semantic Scholar", paper.get("s2_url"), paper.get("s2_citations")


def _sort_papers(papers):
    """年份降序，同年按 GS 被引降序（年份无法解析的排最后）。"""
    def key(p):
        try:
            year = int(str(p.get("year") or "0"))
        except ValueError:
            year = 0
        return (-year, -(p.get("gs_citations") or 0))
    return sorted(papers, key=key)


def _add_kv_table(doc, rows):
    """两列键值表（作者信息用）。"""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v) if v not in (None, "") else "-"
        cells[0].paragraphs[0].runs[0].font.bold = True
    return table


def _add_picture_or_note(doc, path: Path, note_prefix: str):
    """嵌入截图；文件缺失时写一行说明而不是让 docx 生成失败。"""
    if path.exists():
        doc.add_picture(str(path), width=SCREENSHOT_WIDTH)
    else:
        doc.add_paragraph(f"【{note_prefix}缺失：{path.name}】")


def _add_summary_table(doc, papers):
    """近五年论文总表：年内序号 / 标题 / GS 被引 / 核查状态，按年份分隔。

    年份行跨列合并、加粗并放大字号，带该年篇数（如 "2026 年（3 篇）"）；
    年内序号与第三部分逐篇详情的 #N 一致（年内按 GS 被引降序）。
    """
    from collections import Counter
    year_counts = Counter(str(p.get("year") or "未知年份") for p in papers)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(("序号", "论文标题", "GS 被引", "核查状态")):
        hdr[i].text = text
        hdr[i].paragraphs[0].runs[0].font.bold = True

    current_year = None
    year_rank = 0
    for p in papers:  # 调用前已按 年份降序 + 年内 GS 被引降序 排好
        year = str(p.get("year") or "未知年份")
        if year != current_year:
            current_year = year
            year_rank = 0
            row = table.add_row()
            merged = row.cells[0]
            for c in row.cells[1:]:
                merged = merged.merge(c)
            merged.text = f"{year} 年（{year_counts[year]} 篇）"
            run = merged.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(12)
        year_rank += 1
        cells = table.add_row().cells
        cells[0].text = str(year_rank)
        cells[1].text = p.get("title") or "(无标题)"
        cells[2].text = str(p.get("gs_citations") if p.get("gs_citations") is not None else "-")
        cells[3].text = str(p.get("match_status") or "-")
    return table


def generate_report(task_dir: Path) -> Path | None:
    """为单个任务目录生成 <task_id>_report.docx，返回路径；result.json 缺失返回 None。"""
    task_dir = Path(task_dir)
    rj = task_dir / "result.json"
    if not rj.exists():
        return None
    result = json.loads(rj.read_text(encoding="utf-8"))
    task_id = result.get("task_id") or task_dir.name
    shots = task_dir / "screenshots"

    doc = Document()
    _set_style_font(doc.styles["Normal"], BODY_ASCII_FONT, BODY_EASTASIA_FONT, Pt(10.5))
    for _name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        _set_style_font(doc.styles[_name], BODY_ASCII_FONT, HEADING_EASTASIA_FONT)

    # ---- 标题 ----
    doc.add_heading(f"{result.get('person_name') or task_id} · 学者论文核查报告", level=0)
    doc.add_paragraph(
        f"任务编号：{task_id}    生成时间：{datetime.now():%Y-%m-%d %H:%M}    "
        f"任务状态：{result.get('status', '-')}"
        + (f"（{result['note']}）" if result.get("note") else ""))

    # ---- 一、作者信息 ----
    doc.add_heading("一、作者信息", level=1)
    _add_kv_table(doc, [
        ("姓名", result.get("person_name")),
        ("单位", result.get("affiliation")),
        ("研究兴趣", "、".join(result.get("interests") or [])),
        ("总被引", result.get("total_citations")),
        ("h-index", result.get("h_index")),
        ("i10-index", result.get("i10_index")),
        ("Google Scholar 主页", result.get("profile_url")),
    ])
    doc.add_paragraph()
    _add_picture_or_note(doc, shots / f"{task_id}_profile.png", "作者主页截图")

    # ---- 二、近五年论文总表 ----
    papers = _sort_papers(result.get("recent_papers") or [])
    doc.add_heading(f"二、近五年论文总表（共 {len(papers)} 篇）", level=1)
    _add_summary_table(doc, papers)

    # ---- 三、逐篇核查详情（按年份分小节） ----
    doc.add_heading("三、逐篇核查详情", level=1)

    current_year = None
    year_rank = 0  # 年内序号（从 1 开始，年内按 GS 被引降序；全局被引排名见详情表）
    for p in papers:
        year = str(p.get("year") or "未知年份")
        if year != current_year:
            current_year = year
            year_rank = 0
            doc.add_heading(f"{year} 年", level=2)
        year_rank += 1

        src_name, src_url, src_cites = _ext_fields(p)
        doc.add_heading(f"#{year_rank}  {p.get('title') or '(无标题)'}", level=3)
        _add_kv_table(doc, [
            ("年内序号", year_rank),
            ("全局被引排名", p.get("rank")),
            ("发表年份", year),
            ("Google Scholar 被引", p.get("gs_citations")),
            ("核查状态", p.get("match_status")),
            (f"{src_name} 被引", src_cites),
            (f"{src_name} 链接", src_url),
            ("DOI", p.get("doi")),
            ("期刊 / 来源", p.get("journal")),
            ("备注", p.get("note")),
        ])
        doc.add_paragraph("截图证据：")
        shot_name = p.get("screenshot")
        if shot_name:
            _add_picture_or_note(doc, shots / shot_name, "论文截图")
        else:
            doc.add_paragraph("【本篇无截图】")

    out = task_dir / f"{task_id}_report.docx"
    doc.save(str(out))
    return out


def main():
    parser = argparse.ArgumentParser(description="把采集结果导出为 Word 报告")
    parser.add_argument("--task", action="append", help="task_id（可多次指定）")
    parser.add_argument("--all", action="store_true", help="data/ 下所有有 result.json 的任务")
    args = parser.parse_args()

    if not args.task and not args.all:
        parser.error("请指定 --task <task_id> 或 --all")

    task_ids = list(args.task or [])
    if args.all:
        task_ids += sorted(p.parent.name for p in DATA_DIR.glob("*/result.json"))

    ok, skip = 0, 0
    for tid in dict.fromkeys(task_ids):  # 去重保序
        out = generate_report(DATA_DIR / tid)
        if out:
            print(f"[OK] {out}")
            ok += 1
        else:
            print(f"[跳过] {tid}：没有 result.json")
            skip += 1
    print(f"完成：生成 {ok} 份，跳过 {skip} 个")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
